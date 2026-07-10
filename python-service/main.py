"""
QMS Forms - Python Extraction Service
Nhận file .doc/.docx → trả JSON cấu trúc thô + field gợi ý (normalized).
File .doc (định dạng cũ) được convert sang .docx bằng LibreOffice headless trước.
Không có AI — chỉ đọc đúng những gì có trong file + luật chuẩn hóa.
"""

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse, Response
import tempfile
import subprocess
import os
import re
import json

from docx import Document
from extractor import extract_docx_structure
from structure_normalizer import normalize_structure

_FILL_DOTS = re.compile(r"[.…_]{2,}")

app = FastAPI(
    title="QMS Docx Extraction Service",
    description="Trích cấu trúc + gợi ý field từ file .doc/.docx cho QMS Forms",
    version="1.1.0",
)


@app.get("/health")
def health_check():
    return {"status": "ok"}


def _convert_to(src_path: str, target: str) -> str:
    """Convert file sang định dạng target (docx/pdf) bằng LibreOffice. Trả path kết quả."""
    out_dir = tempfile.mkdtemp()
    profile = f"-env:UserInstallation=file:///tmp/lo_{os.getpid()}_{os.path.basename(src_path)}"
    subprocess.run(
        ["soffice", "--headless", profile, "--convert-to", target, "--outdir", out_dir, src_path],
        capture_output=True, timeout=120,
    )
    out = os.path.join(out_dir, os.path.splitext(os.path.basename(src_path))[0] + "." + target)
    if not os.path.exists(out):
        raise RuntimeError(f"Convert sang {target} thất bại")
    return out


def _set_para_text(para, value: str) -> None:
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = value
    else:
        para.add_run(value)


def _is_dots_only(s: str) -> bool:
    s = s.strip()
    return bool(s) and bool(_FILL_DOTS.match(s)) and len(re.sub(r"[.…_\s]", "", s)) == 0


def _fill_paragraph_inline(para, fills: list[tuple[str, str]]) -> bool:
    """Nhãn + dấu chấm CÙNG một đoạn: 'Tên: …………' -> 'Tên: value'."""
    full = para.text
    if not full.strip():
        return False
    new = full
    changed = False
    for label, value in fills:
        if not value or not label or label not in new.lower():
            continue
        idx = new.lower().find(label)
        after = new[idx + len(label):]
        m = _FILL_DOTS.search(after)
        if m:
            # có chuỗi chấm phía sau -> thay bằng giá trị
            new = new[:idx + len(label)] + after[:m.start()] + " " + value + " " + after[m.end():]
            changed = True
        elif after.strip() in ("", ":") and new.rstrip().endswith(":"):
            # dạng 'Nhãn:' (điền ngay sau dấu hai chấm cùng ô)
            new = new.rstrip() + " " + value
            changed = True
    if changed:
        _set_para_text(para, new)
    return changed


def _fill_empty_para(cell, value: str) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    _set_para_text(p, value)


def _fill_doc(path: str, fills: list[tuple[str, str]]) -> None:
    doc = Document(path)

    # 1) Đoạn văn ngoài bảng: 'Nhãn: …………'
    for para in doc.paragraphs:
        _fill_paragraph_inline(para, fills)

    # 2) Bảng — đặt GIÁ TRỊ vào ô trống tương ứng (label cố định sẵn)
    for table in doc.tables:
        rows = table.rows
        n = len(rows)
        for r in range(n):
            cells = rows[r].cells
            seen = set()
            for c, cell in enumerate(cells):
                cid = id(cell._tc)
                if cid in seen:                 # bỏ ô merge nhân bản
                    continue
                seen.add(cid)

                low = cell.text.lower()
                matched = [(l, v) for l, v in fills if v and l and l in low]
                if not matched:
                    continue
                value = matched[0][1]

                # a) trong ô có đoạn toàn dấu chấm -> điền vào đó
                dot_p = next((p for p in cell.paragraphs if _is_dots_only(p.text)), None)
                if dot_p is not None:
                    _set_para_text(dot_p, value)
                    continue
                # b) 'Nhãn: ……' cùng đoạn
                if any(_fill_paragraph_inline(p, matched) for p in cell.paragraphs):
                    continue
                # c) tiêu đề cột -> điền vào ô TRỐNG ngay dưới (cùng cột)
                if r + 1 < n and c < len(rows[r + 1].cells):
                    below = rows[r + 1].cells[c]
                    if id(below._tc) != cid and below.text.strip() == "":
                        _fill_empty_para(below, value)
                        continue

    doc.save(path)


@app.post("/preview-filled")
async def preview_filled(file: UploadFile = File(...), fills: str = Form("[]")):
    """Nhận .docx + danh sách {label,value} → điền vào chỗ trống → trả PDF."""
    name = (file.filename or "").lower()
    if not (name.endswith(".docx") or name.endswith(".doc")):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ .doc/.docx")
    suffix = ".doc" if (name.endswith(".doc") and not name.endswith(".docx")) else ".docx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        # .doc -> .docx trước khi điền
        work = tmp_path
        if suffix == ".doc":
            work = _convert_to(tmp_path, "docx")
        pairs = []
        for item in json.loads(fills or "[]"):
            lbl = (item.get("label") or "").strip().lower()
            val = str(item.get("value") or "").strip()
            if lbl and val:
                pairs.append((lbl, val))
        # điền nhãn dài trước (khớp cụ thể hơn)
        pairs.sort(key=lambda p: -len(p[0]))
        _fill_doc(work, pairs)
        pdf_path = _convert_to(work, "pdf")
        return Response(content=open(pdf_path, "rb").read(), media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi điền/PDF: {e}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


@app.post("/preview")
async def preview(file: UploadFile = File(...)):
    """Nhận .doc/.docx → trả PDF (để xem mẫu gốc khi nhập liệu)."""
    name = (file.filename or "").lower()
    if not (name.endswith(".docx") or name.endswith(".doc")):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ .doc/.docx")
    suffix = ".doc" if (name.endswith(".doc") and not name.endswith(".docx")) else ".docx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        pdf_path = _convert_to(tmp_path, "pdf")
        pdf_bytes = open(pdf_path, "rb").read()
        return Response(content=pdf_bytes, media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi tạo PDF: {e}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _convert_doc_to_docx(src_path: str) -> str:
    """Convert .doc -> .docx bằng LibreOffice headless. Trả về đường dẫn .docx."""
    out_dir = tempfile.mkdtemp()
    # Mỗi lần dùng profile riêng để tránh lock khi gọi đồng thời
    profile = f"-env:UserInstallation=file:///tmp/lo_{os.getpid()}_{os.path.basename(src_path)}"
    proc = subprocess.run(
        ["soffice", "--headless", profile, "--convert-to", "docx", "--outdir", out_dir, src_path],
        capture_output=True, timeout=90,
    )
    base = os.path.splitext(os.path.basename(src_path))[0] + ".docx"
    out_path = os.path.join(out_dir, base)
    if not os.path.exists(out_path):
        raise RuntimeError(
            "Convert .doc -> .docx thất bại: " + (proc.stderr.decode(errors="ignore")[:200] or "unknown")
        )
    return out_path


@app.post("/extract")
async def extract(file: UploadFile = File(...)):
    """
    Nhận file .doc hoặc .docx, trả về JSON gồm:
    - paragraphs, tables (cấu trúc thô + merge cells)
    - normalized.suggested_fields (field gợi ý rule-based, không AI)
    """
    name = (file.filename or "").lower()
    if not (name.endswith(".docx") or name.endswith(".doc")):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .doc hoặc .docx")

    is_doc = name.endswith(".doc") and not name.endswith(".docx")
    suffix = ".doc" if is_doc else ".docx"

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    converted_path = None
    try:
        docx_path = tmp_path
        if is_doc:
            converted_path = _convert_doc_to_docx(tmp_path)
            docx_path = converted_path

        result = extract_docx_structure(docx_path)
        result["normalized"] = normalize_structure(result)
        return JSONResponse(content=result)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi khi đọc file: {str(e)}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        if converted_path:
            try:
                os.unlink(converted_path)
            except OSError:
                pass
