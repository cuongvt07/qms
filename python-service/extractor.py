"""
Lõi trích cấu trúc thô từ .docx bằng python-docx.
Đọc chính xác: paragraphs, tables, gridSpan (horizontal merge), vMerge (vertical merge).
"""

from docx import Document
from docx.oxml.ns import qn
from typing import Any


def _get_grid_span(cell) -> int:
    """Trả về số cột mà ô này span (mặc định 1)."""
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is not None:
        grid_span = tc_pr.find(qn("w:gridSpan"))
        if grid_span is not None:
            return int(grid_span.get(qn("w:val"), 1))
    return 1


def _get_v_merge(cell) -> str | None:
    """
    Trả về:
    - "restart" nếu là ô bắt đầu vertical merge
    - "continue" nếu là ô bị merge phía dưới
    - None nếu không có vertical merge
    """
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is not None:
        v_merge = tc_pr.find(qn("w:vMerge"))
        if v_merge is not None:
            val = v_merge.get(qn("w:val"))
            return "restart" if val == "restart" else "continue"
    return None


def _extract_headers(doc) -> list[str]:
    """
    Đọc text ở HEADER và FOOTER của mọi section.
    Nhiều biểu mẫu ISO đặt TIÊU ĐỀ ở header (letterhead) — thân form không có,
    nên phải đọc header để không mất tiêu đề.
    """
    lines: list[str] = []
    seen: set[str] = set()

    def add(txt: str) -> None:
        t = " ".join((txt or "").split())
        if t and t not in seen:
            seen.add(t)
            lines.append(t)

    for section in doc.sections:
        for holder in (section.header, section.first_page_header,
                       section.even_page_header, section.footer):
            if holder is None:
                continue
            try:
                for p in holder.paragraphs:
                    add(p.text)
                for tb in holder.tables:
                    for row in tb.rows:
                        for cell in row.cells:
                            add(cell.text)
            except Exception:
                continue
    return lines


def extract_docx_structure(file_path: str) -> dict[str, Any]:
    doc = Document(file_path)

    # --- Headers / Footers (thường chứa tiêu đề biểu mẫu) ---
    headers = _extract_headers(doc)

    # --- Paragraphs ---
    paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:  # bỏ qua dòng trắng hoàn toàn
            paragraphs.append({
                "index": idx,
                "text": text,
                "style": para.style.name if para.style else None,
            })

    # --- Tables ---
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows_count = len(table.rows)
        cols_count = len(table.columns) if table.columns else 0

        cells_data = []
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                grid_span = _get_grid_span(cell)
                v_merge = _get_v_merge(cell)
                cell_text = cell.text.strip()

                cells_data.append({
                    "row": r_idx,
                    "col": c_idx,
                    "text": cell_text,
                    "grid_span": grid_span,
                    "v_merge": v_merge,
                })

        tables.append({
            "index": t_idx,
            "rows": rows_count,
            "cols": cols_count,
            "cells": cells_data,
        })

    return {
        "headers": headers,
        "paragraphs": paragraphs,
        "tables": tables,
    }
